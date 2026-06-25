# pyrefly: ignore [missing-import]
import docx
import re
import io
import json
from typing import Dict, Any, List

class JDExtractor:
    def __init__(self, file_path_or_stream: Any):
        self.file_path_or_stream = file_path_or_stream
        self.doc = docx.Document(file_path_or_stream)

    def parse(self) -> Dict[str, Any]:
        """
        Parses the docx file deterministically using a pure State-Machine approach.
        Does not use any LLM or external API calls.
        """
        # Helper to clean up consecutive whitespaces
        def clean_whitespace(text: str) -> str:
            return re.sub(r'\s+', ' ', text).strip()

        # Extract table rows as lines formatted as "Key: Value" or text
        table_lines: List[Dict[str, str]] = []
        for table in self.doc.tables:
            for row in table.rows:
                # Deduplicate cells to avoid repeating merged cell texts
                row_cells = []
                for cell in row.cells:
                    cell_txt = clean_whitespace(cell.text)
                    if cell_txt and (not row_cells or cell_txt != row_cells[-1]):
                        row_cells.append(cell_txt)
                if len(row_cells) == 1:
                    table_lines.append({"text": row_cells[0], "style": "Normal"})
                elif len(row_cells) >= 2:
                    key = row_cells[0]
                    val = " ".join(row_cells[1:])
                    if not key.endswith(':'):
                        table_lines.append({"text": f"{key}: {val}", "style": "Normal"})
                    else:
                        table_lines.append({"text": f"{key} {val}", "style": "Normal"})

        # Extract paragraph texts with style metadata
        paragraph_lines: List[Dict[str, str]] = []
        for p in self.doc.paragraphs:
            text = clean_whitespace(p.text)
            if text:
                paragraph_lines.append({"text": text, "style": p.style.name if p.style else "Normal"})

        # Combine table lines followed by paragraph lines
        clean_lines = table_lines + paragraph_lines

        # Target dictionary mapping to database schema boundaries
        extracted = {
            "job_title": "Unknown",
            "department": "Unknown",
            "experience_required": "Unknown",
            "core_technical_skills": [],
            "behavioral_competencies": [],
            "detailed_responsibilities": ""
        }

        # Step 2: Metadata capture and Headline parsing
        # Isolate Job Title from first 3 non-empty paragraphs/lines
        for line_item in clean_lines[:3]:
            line = line_item["text"]
            m_title = re.match(r'^(?:Job Description|Role|Job Title|Title)\s*:\s*(.*)', line, re.IGNORECASE)
            if m_title:
                extracted["job_title"] = clean_whitespace(m_title.group(1))
                break
        
        # If job title is still Unknown, look at the very first paragraph/line as fallback
        if extracted["job_title"] == "Unknown" and clean_lines:
            first_line = clean_lines[0]["text"]
            if ":" in first_line:
                parts = first_line.split(":", 1)
                extracted["job_title"] = clean_whitespace(parts[1])
            else:
                extracted["job_title"] = clean_whitespace(first_line)

        # Isolate Department: search all text lines for Company: or Department: / Dept:
        for line_item in clean_lines:
            line = line_item["text"]
            m_dept = re.search(r'\b(?:Company|Department|Dept)\s*:\s*(.*)', line, re.IGNORECASE)
            if m_dept:
                val = m_dept.group(1).strip()
                paren_match = re.search(r'\((.*?)\)', val)
                dept_val = clean_whitespace(paren_match.group(1) if paren_match else val)
                
                # If we encounter a "Department/Dept" keyword, it takes precedence and overrides
                is_dept_keyword = "department" in line.lower() or "dept" in line.lower()
                if extracted["department"] == "Unknown" or is_dept_keyword:
                    extracted["department"] = dept_val

        # Step 3 & 4 & 5: State Machine logic
        current_state = None
        allow_normal_in_behavioral = False
        resp_lines: List[str] = []
        exp_lines: List[str] = []

        # List item prefix pattern (bullet formats or digit followed by period)
        list_item_pattern = re.compile(r'^(?:[-*•]|\d+\.)\s*(.*)')

        for item in clean_lines:
            line = item["text"]
            style = item["style"]
            line_lower = line.lower()
            style_lower = style.lower()

            # Determine if the style is a Heading
            is_heading_style = style_lower.startswith("heading") or style_lower in ("title", "subtitle")

            # Determine State Transition / Landmark detection
            next_state = None
            is_landmark = False

            # Landmark A
            if len(line) < 80 and any(term in line_lower for term in ["things you absolutely need", "requirements", "core technical skills", "technical skills", "skills inventory", "things we'd like you to have"]):
                next_state = "TECH_SKILLS"
                is_landmark = True
            # Landmark B
            elif len(line) < 80 and any(term in line_lower for term in ["the vibe check", "culture-fit", "culture fit", "behavioral competencies", "behavioral", "soft skills", "things we explicitly do not want", "things we explicitly do want"]):
                next_state = "BEHAVIORAL"
                is_landmark = True
            # Landmark C
            elif len(line) < 80 and any(term in line_lower for term in ["what you'd actually be doing", "what you’d actually be doing", "responsibilities", "key responsibilities"]):
                next_state = "RESPONSIBILITIES"
                is_landmark = True
            # Landmark D
            elif len(line) < 80 and any(term in line_lower for term in ["what we mean by", "experience required", "experience level"]):
                next_state = "EXPERIENCE"
                is_landmark = True

            # If it's a new structural landmark section, transition state immediately
            if is_landmark:
                current_state = next_state
                # If we transition to EXPERIENCE, check if the landmark line itself has inline value
                if current_state == "EXPERIENCE":
                    m_exp = re.match(r'^(?:Experience Required|Experience Level)\s*:\s*(.*)', line, re.IGNORECASE)
                    if m_exp and m_exp.group(1).strip():
                        exp_lines.append(clean_whitespace(m_exp.group(1)))
                        current_state = None  # Reset state immediately since it's fully matched inline
                # If we transition to BEHAVIORAL, determine if we allow normal paragraphs based on context
                if current_state == "BEHAVIORAL":
                    allow_normal_in_behavioral = any(term in line_lower for term in ["the vibe check", "culture-fit", "culture fit", "behavioral competencies", "behavioral", "soft skills"])
                continue

            # If it is a heading style and NOT a known landmark, reset the state to prevent bleeding
            if is_heading_style:
                current_state = None
                continue

            # Check if this line is a heading/separator that isn't a known landmark, and break current state if so
            # E.g. "on location", "final note", etc.
            if any(term in line_lower for term in ["on location", "how to read between the lines", "final note"]):
                current_state = None
                continue

            # Process line content under the current state
            if current_state == "TECH_SKILLS":
                m_list = list_item_pattern.match(line)
                is_bullet = m_list is not None or "bullet" in style_lower or "list" in style_lower
                if is_bullet:
                    val = clean_whitespace(m_list.group(1) if m_list else line)
                    if val:
                        extracted["core_technical_skills"].append(val)

            elif current_state == "BEHAVIORAL":
                m_list = list_item_pattern.match(line)
                is_bullet = m_list is not None or "bullet" in style_lower or "list" in style_lower or (allow_normal_in_behavioral and style_lower == "normal")
                if is_bullet:
                    val = clean_whitespace(m_list.group(1) if m_list else line)
                    if val:
                        extracted["behavioral_competencies"].append(val)

            elif current_state == "RESPONSIBILITIES":
                resp_lines.append(line)

            elif current_state == "EXPERIENCE":
                exp_lines.append(line)

        # Post-process collections to format output correctly
        if resp_lines:
            extracted["detailed_responsibilities"] = "\n".join(resp_lines)
        else:
            extracted["detailed_responsibilities"] = ""

        if exp_lines:
            extracted["experience_required"] = "\n".join(exp_lines)
        else:
            extracted["experience_required"] = "Unknown"

        return extracted
