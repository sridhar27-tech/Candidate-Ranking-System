import os
import sys
import tempfile
import json

import docx
from .pipeline import run_pipeline
from .schemas import JobDescriptionSchema

def create_dummy_docx(file_path: str):
    """
    Creates a mock docx file with both paragraphs and a structural table
    to comprehensively verify parser capability.
    """
    doc = docx.Document()
    
    # 1. Main Info Table
    doc.add_heading("Job Specification Overview", level=1)
    table = doc.add_table(rows=3, cols=2)
    
    table.cell(0, 0).text = "Job Title"
    table.cell(0, 1).text = "Lead Systems Engineer"
    
    table.cell(1, 0).text = "Department"
    table.cell(1, 1).text = "DevOps Infrastructure"
    
    table.cell(2, 0).text = "Experience Level"
    table.cell(2, 1).text = "8+ years in AWS/GCP architecture"
    
    # 2. Text sections for lists and descriptions
    doc.add_heading("Core Technical Skills", level=2)
    doc.add_paragraph("Candidates must demonstrate competency in:")
    doc.add_paragraph("- Kubernetes and Terraform")
    doc.add_paragraph("- Python, Bash, and Go programming")
    doc.add_paragraph("- CI/CD systems (GitHub Actions/GitLab CI)")
    
    doc.add_heading("Behavioral Competencies", level=2)
    doc.add_paragraph("Candidates will be evaluated on:")
    doc.add_paragraph("- Mentorship of junior engineering staff")
    doc.add_paragraph("- Cross-functional collaboration and design leadership")
    doc.add_paragraph("- High pressure situation management and root cause analysis")
    
    doc.add_heading("Detailed Responsibilities", level=2)
    doc.add_paragraph("In this role, you will:")
    doc.add_paragraph("Ensure high availability and cost optimization of all production clusters.")
    doc.add_paragraph("Architect internal developer platforms to increase deployment speeds.")
    doc.add_paragraph("Lead technical response during high-priority system incidents.")
    
    doc.save(file_path)

def main():
    # Setup temporary file path
    temp_dir = tempfile.gettempdir()
    temp_docx_path = os.path.join(temp_dir, "test_ingestion_doc.docx")
    
    print(f"Creating mock Job Description at: {temp_docx_path}")
    create_dummy_docx(temp_docx_path)
    
    try:
        print("Starting sandbox ingestion pipeline...")
        json_output = run_pipeline(temp_docx_path)
        
        print("\n=== PIPELINE SERIALIZED JSON OUTPUT ===")
        print(json_output)
        print("========================================\n")
        
        # Load and parse output schema
        data = json.loads(json_output)
        model = JobDescriptionSchema(**data)
        
        # Core Assertions
        assert model.job_title == "Lead Systems Engineer", f"Unexpected title: {model.job_title}"
        assert model.department == "DevOps Infrastructure", f"Unexpected department: {model.department}"
        assert "AWS/GCP" in model.experience_required, f"Unexpected experience: {model.experience_required}"
        
        # Technical skills assertion
        assert any("Kubernetes" in skill for skill in model.core_technical_skills), "Kubernetes skill not found"
        assert any("Python" in skill for skill in model.core_technical_skills), "Python skill not found"
        
        # Behavioral competencies assertion
        assert any("Mentorship" in comp for comp in model.behavioral_competencies), "Mentorship competency not found"
        
        # Detailed responsibilities assertion
        assert "Ensure high availability" in model.detailed_responsibilities, "Responsibilities snippet not found"
        
        print("Success: Sandbox validation checks passed successfully!")
        return 0
    except Exception as e:
        print(f"Error during sandbox verification: {e}", file=sys.stderr)
        return 1
    finally:
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
            print("Cleanup: Temporary docx file deleted.")

if __name__ == "__main__":
    sys.exit(main())
